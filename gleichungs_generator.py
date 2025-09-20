#!/usr/bin/env python3
"""Gleichungs-Generator für lineare Gleichungen mit einer Variablen.

Generiert Aufgaben in 5 Schwierigkeitsleveln mit Lösungen als Brüche.
Erstellt Arbeits- und Lösungsblatt im DOCX-Format.
"""

from __future__ import annotations

import argparse
import math
import random
import re
from dataclasses import dataclass, field
from fractions import Fraction

import sympy as sp
from docx import Document
from docx.shared import Pt
from sympy.polys.polyerrors import PolynomialError

x = sp.Symbol("x")

# ---------------------------------------------------------------------------
# Konfiguration
# ---------------------------------------------------------------------------

MAX_ABS = 120
MAX_RESAMPLES = 200

DEFAULT_CONFIG: dict = {
    "seed": 12345,
    "counts": {"L1": 4, "L2": 4, "L3": 6, "L4": 4, "L5": 5},
    "coeff_range": (-12, 12),
    "denom_range": (2, 12),
    "prefer_fraction": 0.85,
    "visual_complexity": "mixed",  # "clean" oder "mixed"
    "solutions": {"improper_and_mixed": True},
    "backend": "docx",
}

# ---------------------------------------------------------------------------
# Hilfsfunktionen
# ---------------------------------------------------------------------------


def nice_frac(fr: Fraction) -> str:
    """Formatiert einen Bruch als String."""
    if fr.denominator == 1:
        return str(fr.numerator)
    return f"{fr.numerator}/{fr.denominator}"


def fraction_to_mixed(fr: Fraction) -> str:
    """Konvertiert einen Bruch in eine gemischte Zahl."""

    if fr.denominator == 1:
        return str(fr.numerator)

    sign = "-" if fr < 0 else ""
    whole, remainder = divmod(abs(fr.numerator), fr.denominator)

    if remainder == 0:
        return f"{sign}{whole}"

    if whole == 0:
        return nice_frac(fr)

    return f"{sign}{whole} {remainder}/{fr.denominator}"


def sample_solution(rng: random.Random, cfg: dict) -> Fraction:
    """Zieht eine Ziellösung, bevorzugt Brüche."""
    a, b = cfg["coeff_range"]
    dmin, dmax = cfg["denom_range"]
    pref = cfg["prefer_fraction"]

    attempts = 0
    while attempts < 50:
        p = rng.randint(a, b)
        if p == 0:
            attempts += 1
            continue
        q = rng.randint(dmin, dmax)
        r = Fraction(p, q)

        # Bevorzuge nicht-ganzzahlige Lösungen
        if r.denominator == 1 and rng.random() < pref:
            attempts += 1
            continue

        # Prüfe ob Lösung im erlaubten Bereich
        if abs(r.numerator) <= MAX_ABS and abs(r.denominator) <= MAX_ABS:
            return r
        attempts += 1

    # Fallback: irgendeine gültige Lösung
    return Fraction(rng.randint(1, 12), rng.randint(2, 6))


def sympy_to_text(expr: sp.Expr, visual_complexity: str = "mixed") -> str:
    """Konvertiert einen SymPy-Ausdruck in lesbaren Text."""

    text = sp.sstr(expr)

    if visual_complexity == "mixed":
        # Entferne Faktoren ±1 vor Symbolen/Klammern.
        text = re.sub(r"(?<![\w])1\*(?=[A-Za-z(])", "", text)
        text = re.sub(r"(?<![\w])-1\*(?=[A-Za-z(])", "-", text)

        # Implizites Mal nur dort, wo es mathematisch eindeutig ist.
        text = re.sub(r"(?<=\d)\*(?=[A-Za-z(])", "", text)
        text = re.sub(r"(?<=\))\*(?=[A-Za-z(])", "", text)
        text = re.sub(r"(?<=[A-Za-z])\*(?=\()", "", text)
    else:
        text = text.replace("*", "·")

    text = re.sub(r"\s+", " ", text)
    text = text.replace("+ -", "- ")
    text = text.replace("- -", "+ ")

    text = re.sub(r"Rational\((-?\d+),\s*(\d+)\)", r"\1/\2", text)

    return text.strip()


def _fraction_ok(fr: Fraction) -> bool:
    """Prüft ob Bruch innerhalb der Zahlengrenzen liegt."""
    return abs(fr.numerator) <= MAX_ABS and abs(fr.denominator) <= MAX_ABS


def _expr_ok(expr: sp.Expr) -> bool:
    """Prüft ob alle Zahlen im Ausdruck innerhalb der Grenzen liegen."""
    for atom in expr.atoms(sp.Rational):
        fr = Fraction(int(atom.p), int(atom.q))
        if not _fraction_ok(fr):
            return False
    return True


def _sympy_to_fraction(value: sp.Rational | int) -> Fraction:
    """Konvertiert einen SymPy-Wert in ``Fraction``."""

    rational = sp.Rational(value)
    return Fraction(int(rational.p), int(rational.q))


def _coefficients_within_limits(expr: sp.Expr) -> bool:
    """Prüft, ob alle Koeffizienten eines Polynoms im Limit liegen."""

    if expr == 0:
        return True

    try:
        poly = sp.Poly(expr, x)
    except PolynomialError:
        return False

    for coeff in poly.all_coeffs():
        if coeff.free_symbols:
            return False

        rationals = coeff.atoms(sp.Rational)
        if not rationals:
            continue

        for atom in rationals:
            if not _fraction_ok(_sympy_to_fraction(atom)):
                return False
    return True


# ---------------------------------------------------------------------------
# Datenklassen
# ---------------------------------------------------------------------------


@dataclass
class Equation:
    """Repräsentiert eine Gleichung."""

    level: str
    sympy_eq: sp.Eq
    text: str
    solution: Fraction
    excluded: set[Fraction] = field(default_factory=set)
    template: str = ""
    params: dict = field(default_factory=dict)


@dataclass
class SolveStep:
    """Ein Schritt im Lösungsweg."""

    description_de: str
    lhs: sp.Expr
    rhs: sp.Expr


@dataclass
class Problem:
    """Aufgabe mit Lösungsschritten."""

    equation: Equation
    steps: list[SolveStep]
    resamples: int = 0
    dops: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Numerische Validierung
# ---------------------------------------------------------------------------


def numeric_limits_ok(eq: Equation) -> bool:
    """Prüft ob alle Zahlen in der Gleichung innerhalb der Limits sind."""
    # Prüfe Ausgangsgleichung
    if not _expr_ok(eq.sympy_eq.lhs) or not _expr_ok(eq.sympy_eq.rhs):
        return False

    # Prüfe nach Expansion
    lhs_exp = sp.expand(eq.sympy_eq.lhs)
    rhs_exp = sp.expand(eq.sympy_eq.rhs)
    if not _expr_ok(lhs_exp) or not _expr_ok(rhs_exp):
        return False

    expr = sp.together(eq.sympy_eq.lhs - eq.sympy_eq.rhs)

    denominators: list[int] = []
    for term in sp.Add.make_args(expr):
        _, denom = sp.fraction(term)
        if denom == 1:
            continue

        coeff, rest = sp.factor(denom).as_coeff_Mul()
        coeff = sp.Rational(coeff)
        num_coeff, den_coeff = coeff.as_numer_denom()

        if abs(num_coeff) != 1:
            denominators.append(abs(int(num_coeff)))
        if den_coeff != 1:
            denominators.append(int(den_coeff))

        if not rest.free_symbols and rest != 1:
            try:
                denominators.append(abs(int(rest)))
            except Exception:  # pragma: no cover - sollte nicht passieren
                pass

    num, den = sp.fraction(expr)
    num = sp.expand(num)
    den = sp.expand(den)

    num_num, num_den = sp.fraction(num)
    if num_den != 1:
        try:
            denominators.append(abs(int(num_den)))
        except Exception:  # pragma: no cover - sollte nicht passieren
            pass
        num = sp.expand(num_num)
        den = sp.expand(den * num_den)

    den_num, den_den = sp.fraction(den)
    if den_den != 1:
        try:
            denominators.append(abs(int(den_den)))
        except Exception:  # pragma: no cover - sollte nicht passieren
            pass
        den = sp.expand(den_num)

    if not _coefficients_within_limits(num):
        return False
    if den != 1 and not _coefficients_within_limits(den):
        return False

    if denominators:
        lcm_val = denominators[0]
        for value in denominators[1:]:
            lcm_val = math.lcm(lcm_val, value)
        if lcm_val > MAX_ABS:
            return False

    return True


# ---------------------------------------------------------------------------
# Beautifier für Textdarstellung
# ---------------------------------------------------------------------------


def beautify_equation(
    lhs: sp.Expr, rhs: sp.Expr, visual_complexity: str = "mixed"
) -> str:
    """Formatiert eine Gleichung schön für die Ausgabe."""
    lhs_str = sympy_to_text(lhs, visual_complexity)
    rhs_str = sympy_to_text(rhs, visual_complexity)

    # Spezielle Formatierungen für "mixed" Komplexität
    if visual_complexity == "mixed":
        # Ersetze 1*(...) durch (...)
        lhs_str = re.sub(r"\b1\s*\(", "(", lhs_str)
        rhs_str = re.sub(r"\b1\s*\(", "(", rhs_str)

        # Ersetze -1*(...) durch -(...)
        lhs_str = re.sub(r"-1\s*\(", "-(", lhs_str)
        rhs_str = re.sub(r"-1\s*\(", "-(", rhs_str)

    return f"{lhs_str} = {rhs_str}"


# ---------------------------------------------------------------------------
# Builder für Level 1: x auf einer Seite
# ---------------------------------------------------------------------------


def build_L1(rng: random.Random, sol: Fraction, cfg: dict) -> Equation:
    """Erstellt Level 1 Gleichung: x nur auf einer Seite."""
    mn, mx = cfg["coeff_range"]

    # Wähle Form
    form = rng.choice(["ax_plus_b", "a_minus_bx", "c_equals_dx_plus_e"])

    if form == "ax_plus_b":
        # ax + b = c
        a = rng.choice([i for i in range(mn, mx + 1) if i != 0])
        b = rng.randint(mn, mx)
        c = a * sol + b

        if not _fraction_ok(Fraction(c)):
            # Retry mit kleineren Werten
            a = rng.choice([i for i in range(-6, 7) if i != 0])
            b = rng.randint(-6, 6)
            c = a * sol + b

        lhs = a * x + b
        rhs = c
        text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))

    elif form == "a_minus_bx":
        # a - bx = c
        b = rng.choice([i for i in range(mn, mx + 1) if i != 0])
        a = rng.randint(mn, mx)
        c = a - b * sol

        if not _fraction_ok(Fraction(c)):
            b = rng.choice([i for i in range(-6, 7) if i != 0])
            a = rng.randint(-6, 6)
            c = a - b * sol

        lhs = a - b * x
        rhs = c
        text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))

    else:  # c_equals_dx_plus_e
        # c = dx + e
        d = rng.choice([i for i in range(mn, mx + 1) if i != 0])
        e = rng.randint(mn, mx)
        c = d * sol + e

        if not _fraction_ok(Fraction(c)):
            d = rng.choice([i for i in range(-6, 7) if i != 0])
            e = rng.randint(-6, 6)
            c = d * sol + e

        lhs = c
        rhs = d * x + e
        text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))

    return Equation("L1", sp.Eq(lhs, rhs, evaluate=False), text, sol)


# ---------------------------------------------------------------------------
# Builder für Level 2: x auf beiden Seiten
# ---------------------------------------------------------------------------


def build_L2(rng: random.Random, sol: Fraction, cfg: dict) -> Equation:
    """Erstellt Level 2 Gleichung: x auf beiden Seiten."""
    mn, mx = cfg["coeff_range"]

    attempts = 0
    while attempts < 50:
        a = rng.choice([i for i in range(mn, mx + 1) if i != 0])
        d = rng.choice([i for i in range(mn, mx + 1) if i != 0 and i != a])
        e = rng.randint(mn, mx)
        b = d * sol + e - a * sol

        if _fraction_ok(Fraction(b)):
            lhs = a * x + b
            rhs = d * x + e
            text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
            return Equation("L2", sp.Eq(lhs, rhs, evaluate=False), text, sol)
        attempts += 1

    # Fallback
    lhs = 2 * x + 3
    rhs = x + 3 + sol
    text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
    return Equation("L2", sp.Eq(lhs, rhs, evaluate=False), text, sol)


# ---------------------------------------------------------------------------
# Builder für Level 3: Eine Klammer (je 2× +, -, ×)
# ---------------------------------------------------------------------------


def build_L3(
    rng: random.Random, sol: Fraction, cfg: dict, pattern_index: int = None
) -> Equation:
    """Erstellt Level 3 Gleichung: Eine Klammer.

    pattern_index: 0,1 = plus; 2,3 = minus; 4,5 = times
    """
    mn, mx = -8, 8  # Kleinere Werte für Klammern

    # Bestimme Pattern basierend auf Index
    if pattern_index is None:
        pattern = rng.choice(["plus", "minus", "times"])
    else:
        if pattern_index < 2:
            pattern = "plus"
        elif pattern_index < 4:
            pattern = "minus"
        else:
            pattern = "times"

    m = rng.choice([i for i in range(mn, mx + 1) if i != 0])
    n = rng.randint(mn, mx)
    p = rng.choice([i for i in range(mn, mx + 1) if i != 0])

    if pattern == "plus":
        # k + (mx + n) = px + q
        k = rng.randint(mn, mx)
        q = k + (m * sol + n) - p * sol

        if _fraction_ok(Fraction(q)):
            lhs = k + (m * x + n)
            rhs = p * x + q
            text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
            template = "k + (mx + n) = px + q"
        else:
            # Fallback
            lhs = 2 + (x + 1)
            rhs = 2 * x + 1 + 2 - sol
            text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
            template = "plus"

    elif pattern == "minus":
        # k - (mx + n) = px + q
        k = rng.randint(mn, mx)
        q = k - (m * sol + n) - p * sol

        if _fraction_ok(Fraction(q)):
            lhs = k - (m * x + n)
            rhs = p * x + q
            text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
            template = "k - (mx + n) = px + q"
        else:
            # Fallback
            lhs = 5 - (x + 2)
            rhs = x + 3 - 2 * sol
            text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
            template = "minus"

    else:  # times
        # t(mx + n) = px + q
        t = rng.choice([i for i in range(-6, 7) if i != 0])
        q = t * (m * sol + n) - p * sol

        if _fraction_ok(Fraction(q)):
            lhs = t * (m * x + n)
            rhs = p * x + q

            # Bei "mixed" complexity: implizites Mal
            if cfg.get("visual_complexity", "mixed") == "mixed":
                lhs_str = f"{t}({sympy_to_text(m * x + n, 'mixed')})"
                rhs_str = sympy_to_text(rhs, "mixed")
                text = f"{lhs_str} = {rhs_str}"
            else:
                text = beautify_equation(
                    lhs, rhs, cfg.get("visual_complexity", "clean")
                )
            template = "t(mx + n) = px + q"
        else:
            # Fallback
            lhs = 3 * (x + 1)
            rhs = 4 * x + 3 - sol
            text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
            template = "times"

    return Equation("L3", sp.Eq(lhs, rhs, evaluate=False), text, sol, template=template)


# ---------------------------------------------------------------------------
# Builder für Level 4: Gemischte Klammern
# ---------------------------------------------------------------------------


def build_L4(rng: random.Random, sol: Fraction, cfg: dict) -> Equation:
    """Erstellt Level 4 Gleichung: Gemischte Klammern."""
    # Kleinere Bereiche für komplexere Gleichungen
    attempts = 0
    while attempts < 50:
        a = rng.choice([i for i in range(-4, 5) if i != 0])
        m = rng.choice([i for i in range(-4, 5) if i != 0])
        n = rng.randint(-6, 6)
        b = rng.randint(-6, 6)
        c = rng.choice([i for i in range(-4, 5) if i != 0])
        d = rng.randint(-6, 6)
        p = rng.randint(-6, 6)
        q = rng.choice([i for i in range(-4, 5) if i != 0])
        r = rng.randint(-6, 6)
        s = rng.choice([i for i in range(-3, 4) if i != 0])
        t = rng.choice([i for i in range(-3, 4) if i != 0])

        # Berechne u so dass die Gleichung bei x=sol erfüllt ist
        left_at_sol = a * (m * sol + n) + b - (c * sol + d)
        right_without_u = p - (q * sol + r) + s * t * sol

        if s != 0:
            u = (left_at_sol - right_without_u) / s

            if _fraction_ok(Fraction(u)):
                lhs = a * (m * x + n) + b - (c * x + d)
                rhs = p - (q * x + r) + s * (t * x + u)

                # Formatierung für mixed complexity
                if cfg.get("visual_complexity", "mixed") == "mixed":
                    # Implizites Mal vor Klammern
                    lhs_str = (
                        f"{a}({sympy_to_text(m * x + n, 'mixed')}) + {b} - "
                        f"({sympy_to_text(c * x + d, 'mixed')})"
                    )
                    rhs_str = (
                        f"{p} - ({sympy_to_text(q * x + r, 'mixed')}) + "
                        f"{s}({sympy_to_text(t * x + u, 'mixed')})"
                    )
                    text = f"{lhs_str} = {rhs_str}".replace("+ -", "- ")
                else:
                    text = beautify_equation(
                        lhs, rhs, cfg.get("visual_complexity", "clean")
                    )

                eq = Equation("L4", sp.Eq(lhs, rhs, evaluate=False), text, sol)
                if numeric_limits_ok(eq):
                    return eq
        attempts += 1

    # Fallback
    lhs = 2 * (x + 1) - (x - 1)
    rhs = x + 3
    text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
    return Equation("L4", sp.Eq(lhs, rhs, evaluate=False), text, sol)


# ---------------------------------------------------------------------------
# Builder für Level 5: Bruchgleichungen
# ---------------------------------------------------------------------------


def build_L5(rng: random.Random, sol: Fraction, cfg: dict) -> Equation:
    """Erstellt Level 5 Gleichung: Bruchgleichungen."""
    variant = rng.choice(["simple", "two_fracs", "cross"])

    if variant == "simple":
        # (a/b)x = c
        a = rng.randint(1, 8)
        b = rng.randint(2, 8)
        c = Fraction(a, b) * sol

        if _fraction_ok(c):
            lhs = sp.Rational(a, b) * x
            rhs = (
                sp.Rational(c.numerator, c.denominator)
                if c.denominator != 1
                else c.numerator
            )
            text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
            return Equation("L5", sp.Eq(lhs, rhs, evaluate=False), text, sol)

    elif variant == "two_fracs":
        # (ax + b)/c + (dx + e)/f = g
        attempts = 0
        while attempts < 30:
            a = rng.choice([i for i in range(-4, 5) if i != 0])
            b = rng.randint(-6, 6)
            c = rng.randint(2, 8)
            d = rng.choice([i for i in range(-4, 5) if i != 0])
            e = rng.randint(-6, 6)
            f = rng.randint(2, 8)

            # Berechne g so dass die Gleichung bei x=sol erfüllt ist
            g = Fraction(a * sol + b, c) + Fraction(d * sol + e, f)

            if _fraction_ok(g) and math.lcm(c, f) <= MAX_ABS:
                lhs = (a * x + b) / c + (d * x + e) / f
                rhs = (
                    sp.Rational(g.numerator, g.denominator)
                    if g.denominator != 1
                    else g.numerator
                )
                text = beautify_equation(
                    lhs, rhs, cfg.get("visual_complexity", "mixed")
                )
                excluded = set()
                return Equation(
                    "L5", sp.Eq(lhs, rhs, evaluate=False), text, sol, excluded
                )
            attempts += 1

    else:  # cross
        # (ax + b)/(cx + d) = e/f
        attempts = 0
        while attempts < 30:
            a = rng.choice([i for i in range(-4, 5) if i != 0])
            b = rng.randint(-6, 6)
            c = rng.choice([i for i in range(-3, 4) if i != 0])
            d = rng.randint(-6, 6)

            # Prüfe ob sol ein Ausschlusswert wäre
            if c * sol + d == 0:
                attempts += 1
                continue

            # Berechne e/f
            val = Fraction(a * sol + b, c * sol + d)

            if _fraction_ok(val):
                e = val.numerator
                f = val.denominator

                lhs = (a * x + b) / (c * x + d)
                rhs = sp.Rational(e, f) if f != 1 else e
                text = beautify_equation(
                    lhs, rhs, cfg.get("visual_complexity", "mixed")
                )

                # Ausschlusswert
                excluded = {Fraction(-d, c)} if c != 0 else set()
                return Equation(
                    "L5", sp.Eq(lhs, rhs, evaluate=False), text, sol, excluded
                )
            attempts += 1

    # Fallback: einfache Bruchgleichung
    lhs = sp.Rational(1, 2) * x
    rhs = (
        sp.Rational(sol.numerator, 2 * sol.denominator)
        if sol.denominator != 1
        else sp.Rational(sol.numerator, 2)
    )
    text = beautify_equation(lhs, rhs, cfg.get("visual_complexity", "mixed"))
    return Equation("L5", sp.Eq(lhs, rhs, evaluate=False), text, sol)


# ---------------------------------------------------------------------------
# Validierung
# ---------------------------------------------------------------------------


def validate(eq: Equation) -> bool:
    """Validiert eine Gleichung."""
    try:
        solutions = sp.solve(eq.sympy_eq, x)
    except Exception:
        return False

    if len(solutions) != 1:
        return False

    sol_sympy = sp.nsimplify(solutions[0])
    if not sol_sympy.is_rational:
        return False

    sol_frac = _sympy_to_fraction(sol_sympy)
    if sol_frac != eq.solution:
        return False

    if any(sol_frac == excl for excl in eq.excluded):
        return False

    return True


def canonical_key(eq: Equation) -> str:
    """Erzeugt einen kanonischen Schlüssel für Duplikaterkennung."""
    expr = sp.expand(sp.together(eq.sympy_eq.lhs - eq.sympy_eq.rhs))
    expr = sp.collect(expr, x)
    return f"{expr}|{eq.solution}"


# ---------------------------------------------------------------------------
# Solver mit Schrittverfolgung
# ---------------------------------------------------------------------------


def solve_steps(eq: Equation, cfg: dict) -> list[SolveStep]:
    """Erzeugt Lösungsschritte für eine Gleichung."""
    steps: list[SolveStep] = []

    # Schritt 1: Ausgangsgleichung
    steps.append(SolveStep("Ausgangsgleichung", eq.sympy_eq.lhs, eq.sympy_eq.rhs))

    # Wenn Ausschlusswerte vorhanden, erwähne sie
    if eq.excluded:
        excl_str = ", ".join([f"x ≠ {nice_frac(e)}" for e in sorted(eq.excluded)])
        steps.append(
            SolveStep(f"Definitionsmenge: {excl_str}", eq.sympy_eq.lhs, eq.sympy_eq.rhs)
        )

    # Schritt 2: Klammern auflösen
    lhs = sp.expand(eq.sympy_eq.lhs)
    rhs = sp.expand(eq.sympy_eq.rhs)

    if lhs != eq.sympy_eq.lhs or rhs != eq.sympy_eq.rhs:
        steps.append(SolveStep("Klammern auflösen", lhs, rhs))

    # Schritt 3: Brüche beseitigen (falls vorhanden)
    lhs_together = sp.together(lhs)
    rhs_together = sp.together(rhs)

    lhs_num, lhs_den = sp.fraction(lhs_together)
    rhs_num, rhs_den = sp.fraction(rhs_together)

    denominators = [den for den in (lhs_den, rhs_den) if den != 1]

    if denominators:
        lcm_expr = denominators[0]
        for denom in denominators[1:]:
            lcm_expr = sp.lcm(lcm_expr, denom)

        if _expr_ok(lcm_expr):
            lhs = sp.simplify(lhs * lcm_expr)
            rhs = sp.simplify(rhs * lcm_expr)
            lcm_text = sympy_to_text(lcm_expr, cfg.get("visual_complexity", "mixed"))
            steps.append(
                SolveStep(f"Mit Hauptnenner {lcm_text} multiplizieren", lhs, rhs)
            )

    # Schritt 4: Alle Terme auf eine Seite
    expr = sp.expand(lhs - rhs)
    steps.append(SolveStep("Alle Terme auf eine Seite bringen", expr, sp.Integer(0)))

    # Schritt 5: Nach x zusammenfassen
    expr = sp.collect(expr, x)
    steps.append(SolveStep("Nach x zusammenfassen", expr, sp.Integer(0)))

    # Schritt 6: Koeffizienten extrahieren und GCD-Normalisierung
    if expr.has(x):
        poly = sp.Poly(expr, x, domain="QQ")
        coeffs = poly.all_coeffs()

        if len(coeffs) >= 2:
            a_frac = _sympy_to_fraction(coeffs[0])
            b_frac = _sympy_to_fraction(coeffs[1])

            den_lcm = math.lcm(a_frac.denominator, b_frac.denominator)
            a_int = a_frac.numerator * (den_lcm // a_frac.denominator)
            b_int = b_frac.numerator * (den_lcm // b_frac.denominator)

            expr_int = sp.Integer(a_int) * x + sp.Integer(b_int)

            g = math.gcd(abs(a_int), abs(b_int))
            if g > 1:
                a_int //= g
                b_int //= g
                expr_int = sp.Integer(a_int) * x + sp.Integer(b_int)
                steps.append(
                    SolveStep(
                        f"Durch gemeinsamen Teiler {g} kürzen", expr_int, sp.Integer(0)
                    )
                )

            if a_int != 0:
                sol_frac = Fraction(-b_int, a_int)
                sol_sympy = sp.Rational(sol_frac.numerator, sol_frac.denominator)
                steps.append(SolveStep(f"Durch {a_int} teilen", x, sol_sympy))

                if cfg.get("solutions", {}).get("improper_and_mixed", True):
                    if sol_frac.denominator != 1:
                        mixed = fraction_to_mixed(sol_frac)
                        if mixed != nice_frac(sol_frac):
                            steps.append(
                                SolveStep(f"Als gemischte Zahl: {mixed}", x, sol_sympy)
                            )

    return steps


# ---------------------------------------------------------------------------
# Generator
# ---------------------------------------------------------------------------


def generate_equations(cfg: dict) -> list[Problem]:
    """Generiert alle Gleichungen gemäß Konfiguration."""
    rng = random.Random(cfg["seed"])
    problems: list[Problem] = []
    seen_keys: set[str] = set()
    max_resamples = cfg.get("max_resamples", MAX_RESAMPLES)
    strict = cfg.get("strict_limits", True)

    def try_register(eq: Equation, attempts: int) -> bool:
        if strict and not numeric_limits_ok(eq):
            return False
        if not validate(eq):
            return False

        key = canonical_key(eq)
        if key in seen_keys:
            return False

        seen_keys.add(key)
        steps = solve_steps(eq, cfg)
        problems.append(
            Problem(eq, steps, resamples=attempts, dops=eq.params.get("dops", []))
        )
        return True

    def require_success(level: str) -> None:
        raise RuntimeError(
            f"Keine gültige Gleichung für Level {level} nach {max_resamples} Versuchen."
        )

    for _ in range(cfg["counts"]["L1"]):
        for attempts in range(max_resamples):
            sol = sample_solution(rng, cfg)
            if try_register(build_L1(rng, sol, cfg), attempts):
                break
        else:
            require_success("L1")

    for _ in range(cfg["counts"]["L2"]):
        for attempts in range(max_resamples):
            sol = sample_solution(rng, cfg)
            if try_register(build_L2(rng, sol, cfg), attempts):
                break
        else:
            require_success("L2")

    for idx in range(cfg["counts"]["L3"]):
        for attempts in range(max_resamples):
            sol = sample_solution(rng, cfg)
            if try_register(build_L3(rng, sol, cfg, idx), attempts):
                break
        else:
            require_success("L3")

    for _ in range(cfg["counts"]["L4"]):
        for attempts in range(max_resamples):
            sol = sample_solution(rng, cfg)
            if try_register(build_L4(rng, sol, cfg), attempts):
                break
        else:
            require_success("L4")

    for _ in range(cfg["counts"]["L5"]):
        for attempts in range(max_resamples):
            sol = sample_solution(rng, cfg)
            if try_register(build_L5(rng, sol, cfg), attempts):
                break
        else:
            require_success("L5")

    return problems


# ---------------------------------------------------------------------------
# DOCX Renderer
# ---------------------------------------------------------------------------


def render_arbeitsblatt(problems: list[Problem], filename: str):
    """Erstellt das Arbeitsblatt."""
    doc = Document()
    doc.add_heading("Lineare Gleichungen – Arbeitsblatt", level=1)

    # Gruppiere nach Level
    by_level = {}
    for p in problems:
        level = p.equation.level
        if level not in by_level:
            by_level[level] = []
        by_level[level].append(p)

    # Durchgehende Nummerierung
    num = 1

    for level in ["L1", "L2", "L3", "L4", "L5"]:
        if level not in by_level:
            continue

        doc.add_heading(f"Level {level[1]}", level=2)

        for problem in by_level[level]:
            p = doc.add_paragraph(f"{num}. {problem.equation.text}")
            p.runs[0].font.size = Pt(12)

            # Leerraum zum Rechnen
            for _ in range(4):
                doc.add_paragraph("")

            num += 1

    doc.save(filename)


def render_loesungsblatt(problems: list[Problem], filename: str, cfg: dict):
    """Erstellt das Lösungsblatt."""
    doc = Document()
    doc.add_heading("Lineare Gleichungen – Lösungsblatt", level=1)

    # Gruppiere nach Level
    by_level = {}
    for p in problems:
        level = p.equation.level
        if level not in by_level:
            by_level[level] = []
        by_level[level].append(p)

    # Durchgehende Nummerierung
    num = 1

    for level in ["L1", "L2", "L3", "L4", "L5"]:
        if level not in by_level:
            continue

        doc.add_heading(f"Level {level[1]}", level=2)

        for problem in by_level[level]:
            doc.add_heading(f"Aufgabe {num}: {problem.equation.text}", level=3)

            # Lösungsschritte
            for step in problem.steps:
                if step.description_de.startswith("Definitionsmenge"):
                    p = doc.add_paragraph(f"• {step.description_de}")
                else:
                    eq_text = beautify_equation(
                        step.lhs, step.rhs, cfg.get("visual_complexity", "mixed")
                    )
                    p = doc.add_paragraph(f"• {step.description_de}: {eq_text}")
                p.runs[0].font.size = Pt(11)

            # Finale Lösung
            sol_frac = problem.equation.solution
            sol_str = nice_frac(sol_frac)

            if cfg.get("solutions", {}).get("improper_and_mixed", True):
                mixed = fraction_to_mixed(sol_frac)
                if mixed != sol_str:
                    p = doc.add_paragraph(f"Lösung: x = {sol_str} = {mixed}")
                else:
                    p = doc.add_paragraph(f"Lösung: x = {sol_str}")
            else:
                p = doc.add_paragraph(f"Lösung: x = {sol_str}")

            p.runs[0].bold = True
            doc.add_paragraph("")
            num += 1

    doc.save(filename)


# ---------------------------------------------------------------------------
# Hauptfunktion
# ---------------------------------------------------------------------------


def main() -> None:
    """Hauptfunktion."""
    parser = argparse.ArgumentParser(description="Generator für lineare Gleichungen")
    parser.add_argument("--seed", type=int, default=DEFAULT_CONFIG["seed"])
    parser.add_argument("--visual", choices=["clean", "mixed"], default="mixed")
    args = parser.parse_args()

    cfg = DEFAULT_CONFIG.copy()
    cfg["seed"] = args.seed
    cfg["visual_complexity"] = args.visual

    print(f"Generiere Gleichungen mit Seed {cfg['seed']}...")
    problems = generate_equations(cfg)

    print(f"Erstellt: {len(problems)} Aufgaben")

    render_arbeitsblatt(problems, "arbeitsblatt.docx")
    render_loesungsblatt(problems, "loesungsblatt.docx", cfg)

    print("Dateien erstellt: arbeitsblatt.docx, loesungsblatt.docx")


if __name__ == "__main__":
    main()
