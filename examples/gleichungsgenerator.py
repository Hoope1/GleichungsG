#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Beispielskript zur Generierung linearer Gleichungen.

Dieses Skript bleibt als einfaches, historisches Beispiel erhalten. Die
empfohlene und weiterentwickelte Variante befindet sich in
``gleichungs_generator.py``.

Ausführung::

    python examples/gleichungsgenerator.py
"""
from fractions import Fraction
from dataclasses import dataclass
from random import randint, choice, seed, shuffle
from typing import Tuple, List, Callable, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------------- Konfiguration ----------------------------------

seed(42)  # Für reproduzierbare Ergebnisse; anpassen/auskommentieren nach Bedarf

CONFIG = {
    "counts": {
        "sehr_leicht": 10,
        "leicht": 10,
        "mittel": 10,
        "schwer": 8,
        "brueche": 10,
    },
    "leerzeilen_zwischen_aufgaben": 2,   # Platz zum Rechnen
    "schriftgroesse_punkte": 14,
    "arbeitsblatt_datei": "Arbeitsblatt_Gleichungen.docx",
    "loesungen_datei": "Loesungen_Gleichungen.docx",
    "ziel_mehr_brueche": True,  # versucht, ganzzahlige x möglichst zu vermeiden
}

# ----------------------------- Hilfsfunktionen --------------------------------

def gcd(a: int, b: int) -> int:
    while b:
        a, b = b, a % b
    return abs(a)

def lcm(a: int, b: int) -> int:
    return abs(a*b) // gcd(a, b) if a and b else 0

def lcm_list(vals: List[int]) -> int:
    cur = 1
    for v in vals:
        cur = lcm(cur, abs(v))
    return cur

def rnd_nonzero(a: int, b: int) -> int:
    """Ganzzahl in [a,b], aber nicht 0."""
    while True:
        v = randint(a, b)
        if v != 0:
            return v

def rnd_coeff_avoid_int_solution() -> Fraction:
    """
    Wähle einen Bruchkoeffizienten für x, um ganzzahlige Lösungen seltener zu machen.
    Beispiele: 1/2, 2/3, -3/4, ...
    """
    num = rnd_nonzero(-5, 5)
    den = rnd_nonzero(2, 6)  # vermeidet 1, begünstigt Brüche
    return Fraction(num, den)

def format_frac(fr: Fraction) -> str:
    """Format für eine rationale Zahl (als echter Bruch, ggf. mit Minus)."""
    if fr.denominator == 1:
        return f"{fr.numerator}"
    return f"{fr.numerator}/{fr.denominator}"

def format_coeff_times_x(a: Fraction) -> str:
    """
    Gibt a*x ohne '*' aus:
    - 1*x  -> 'x'
    - -1*x -> '-x'
    - 3/2*x -> '(3/2)x'
    - 3*x -> '3x'
    """
    if a == 1:
        return "x"
    if a == -1:
        return "-x"
    if a.denominator == 1:
        return f"{a.numerator}x"
    else:
        return f"({format_frac(a)})x"

def fmt_term(value: Fraction, with_sign: bool = True) -> str:
    """
    Format konstantes Glied:
    with_sign=True -> inklusive Vorzeichen z.B. '+ 5', '- 3/2'
    """
    s = format_frac(value)
    if with_sign:
        if value >= 0:
            return f"+ {s}"
        else:
            return f"- {format_frac(-value)}"
    else:
        return s

def tidy_spaces(expr: str) -> str:
    return (
        expr.replace("+ -", "- ")
            .replace("- -", "+ ")
            .replace("+  -", "- ")
            .replace("-  -", "+ ")
            .replace("  ", " ")
            .replace("=  ", "= ")
            .replace("  =", " =")
            .strip()
    )

# ---------------------- Algebraischer Kern (lineare Gleichung) ----------------

@dataclass
class LinearEquation:
    # ax + b = cx + d  (alle Koeffizienten rational)
    a: Fraction
    b: Fraction
    c: Fraction
    d: Fraction

    def solution(self) -> Fraction:
        # ax + b = cx + d -> (a-c)x = (d-b) -> x = (d-b)/(a-c)
        if self.a == self.c:
            raise ValueError("Ungültige Gleichung: a == c -> keine eindeutige Lösung.")
        return (self.d - self.b) / (self.a - self.c)

    def as_string(self) -> str:
        left = []
        # ax
        if self.a != 0:
            left.append(format_coeff_times_x(self.a))
        # + b
        if self.b != 0:
            left.append(fmt_term(self.b, with_sign=(len(left) > 0)))
        if not left:
            left = ["0"]
        left_str = " ".join(left)

        right = []
        if self.c != 0:
            right.append(format_coeff_times_x(self.c))
        if self.d != 0:
            right.append(fmt_term(self.d, with_sign=(len(right) > 0)))
        if not right:
            right = ["0"]
        right_str = " ".join(right)

        return tidy_spaces(f"{left_str} = {right_str}")

# ---------------------- Erzeugung “kanonisch” und dann tarnen ----------------

def pick_target_x(prefer_fraction=True) -> Fraction:
    """
    Ziel-Lösung x. Wenn prefer_fraction==True, häufig echte Brüche.
    """
    if prefer_fraction and randint(1, 100) <= 80:
        # echte Brüche in sinnvoller Größe
        num = rnd_nonzero(-12, 12)
        den = rnd_nonzero(2, 12)
        return Fraction(num, den)
    else:
        # kleine ganze Zahlen auch zulassen
        return Fraction(rnd_nonzero(-12, 12), 1)

def build_base_equation(x_val: Fraction) -> LinearEquation:
    """
    Erzeuge eine einfache Form: a*x + b = c*x + d mit kleinen ganzzahligen a,c,b,d
    die genau x_val als Lösung hat.
    """
    # Wähle a-c != 0
    a = Fraction(rnd_nonzero(-6, 6))
    c = a - rnd_nonzero(-6, 6)  # -> a-c kann groß sein, aber ok
    while a == c:
        c = Fraction(rnd_nonzero(-6, 6))

    # d - b = (a-c)*x  -> wähle b zufällig und setze d entsprechend
    b = Fraction(randint(-18, 18))
    d = (a - c) * x_val + b
    return LinearEquation(a, b, c, d)

# ---------------------- Tarn- und Formatierungsfunktionen ---------------------

def to_very_easy_form(eq: LinearEquation) -> LinearEquation:
    """
    Ziel: Muster wie '3x - 7 = 12' oder '5 - 2x = 17' usw.
    Wir sorgen (probabilistisch) dafür, dass eine Seite nur ax + b und die andere nur Zahl ist.
    """
    x = eq.solution()  # korrekt
    # Wir setzen gezielt eine Seite ohne x:
    side = choice(["left_const", "right_const"])
    a = Fraction(rnd_nonzero(-9, 9))
    b = Fraction(randint(-20, 20))
    if side == "left_const":
        # ax + b = konst -> konst = a*x + b
        konst = a * x + b
        return LinearEquation(a, b, Fraction(0), konst)
    else:
        # konst = ax + b
        konst = a * x + b
        return LinearEquation(Fraction(0), konst, a, b)

def to_easy_both_sides(eq: LinearEquation) -> LinearEquation:
    """
    x auf beiden Seiten, aber ohne Klammern.
    """
    x = eq.solution()
    a = Fraction(rnd_nonzero(-9, 9))
    c = a - rnd_nonzero(-6, 6)
    while a == c:
        c = a - rnd_nonzero(-6, 6)
    b = Fraction(randint(-20, 20))
    d = (a - c) * x + b
    return LinearEquation(a, b, c, d)

@dataclass
class DisplayEquation:
    text: str
    # Für die Lösungsschritte behalten wir ein “sauberes” ax+b=cx+d bei:
    canonical: LinearEquation
    kind: str  # Kategorie

def expand(a: Fraction, expr: Tuple[Fraction, Fraction]) -> Tuple[Fraction, Fraction]:
    """Multipliziert a*(p*x + q) -> (a*p, a*q)"""
    p, q = expr
    return (a * p, a * q)

def pretty_parenthesis(ax_b: Tuple[Fraction, Fraction]) -> str:
    p, q = ax_b
    parts = []
    if p != 0:
        parts.append(format_coeff_times_x(p))
    if q != 0:
        parts.append(fmt_term(q, with_sign=(len(parts) > 0)))
    if not parts:
        parts = ["0"]
    return f"({tidy_spaces(' '.join(parts))})"

def make_parenthesis_side(x_val: Fraction,
                          allow_sign: List[str]) -> Tuple[str, Tuple[Fraction, Fraction]]:
    """
    Erzeuge einen Ausdruck in der Form:
      '+ (ax+b)', '- (ax+b)', oder 'k(ax+b)' (Multiplikation),
    und liefere sowohl die Textdarstellung als auch die expandierten (p, q).
    """
    p = Fraction(rnd_nonzero(-5, 5))
    q = Fraction(randint(-10, 10))
    inner = (p, q)
    op = choice(allow_sign)
    if op == "+":
        text = f"+ {pretty_parenthesis(inner)}"
        return (text, inner)
    elif op == "-":
        # -(ax+b) = (-1)*(ax+b)
        text = f"- {pretty_parenthesis(inner)}"
        return (text, (-inner[0], -inner[1]))
    else:
        # Multiplikation k(ax+b) mit k != 0, +-2..6
        k = Fraction(rnd_nonzero(-6, 6))
        inner_exp = expand(k, inner)
        text = f"{k.numerator if k.denominator==1 else '('+format_frac(k)+')'}{pretty_parenthesis(inner)}"
        return (text, inner_exp)

def assemble_sides(
    left_chunks: List[Tuple[str, Tuple[Fraction, Fraction]]],
    right_chunks: List[Tuple[str, Tuple[Fraction, Fraction]]],
) -> Tuple[str, LinearEquation]:
    """
    Summiere alle Chunks (jeder Chunk ist Text und (px+q)), erhalte ax+b=cx+d
    """
    aL = Fraction(0); bL = Fraction(0)
    aR = Fraction(0); bR = Fraction(0)

    left_text_parts = []
    for t, (p, q) in left_chunks:
        left_text_parts.append(t)
        aL += p; bL += q

    right_text_parts = []
    for t, (p, q) in right_chunks:
        right_text_parts.append(t)
        aR += p; bR += q

    left_text = tidy_spaces(" ".join(part if part.startswith(("+","-")) else f"+ {part}" for part in left_text_parts)).lstrip("+ ").strip()
    right_text = tidy_spaces(" ".join(part if part.startswith(("+","-")) else f"+ {part}" for part in right_text_parts)).lstrip("+ ").strip()

    if not left_text:  left_text  = "0"
    if not right_text: right_text = "0"

    return f"{left_text} = {right_text}", LinearEquation(aL, bL, aR, bR)

def to_medium_with_parenthesis(x_val: Fraction) -> DisplayEquation:
    """
    Erzeuge Gleichungen mit genau EINER Klammerseite (plus, minus oder multiplikativ),
    andere Seite ohne Klammer (ax+b).
    """
    # Seite mit Klammer:
    klammerseite = choice(["left", "right"])
    # Erzeuge Klammer mit + / - / *:
    op = choice(["+", "-", "*"])
    textK, (pK, qK) = make_parenthesis_side(x_val, ["+", "-", "*"])

    # andere Seite: einfacher ax+b
    a2 = Fraction(rnd_nonzero(-6, 6))
    b2 = Fraction(randint(-15, 15))

    if klammerseite == "left":
        left_chunks = [(textK, (pK, qK))]
        right_chunks = [("", (a2, b2))]
    else:
        left_chunks = [("", (a2, b2))]
        right_chunks = [(textK, (pK, qK))]

    txt, canon = assemble_sides(left_chunks, right_chunks)
    return DisplayEquation(text=txt, canonical=canon, kind="mittel")

def to_harder_mixed_parenthesis(x_val: Fraction) -> DisplayEquation:
    """
    Gemischte Klammern auf beiden Seiten, mehrere Terme addiert.
    """
    # Zahl der Chunks je Seite:
    nL = randint(1, 2)
    nR = randint(1, 2)
    left_chunks = [make_parenthesis_side(x_val, ["+", "-", "*"]) for _ in range(nL)]
    right_chunks = [make_parenthesis_side(x_val, ["+", "-", "*"]) for _ in range(nR)]

    # Manchmal noch ein nacktes ax+b ohne Klammer einstreuen
    if randint(1, 100) <= 50:
        left_chunks.append(("", (Fraction(rnd_nonzero(-5, 5)), Fraction(randint(-8, 8)))))
    if randint(1, 100) <= 50:
        right_chunks.append(("", (Fraction(rnd_nonzero(-5, 5)), Fraction(randint(-8, 8)))))

    txt, canon = assemble_sides(left_chunks, right_chunks)
    return DisplayEquation(text=txt, canonical=canon, kind="schwer")

def to_fraction_equation(x_val: Fraction) -> DisplayEquation:
    """Erzeuge eine lineare Bruchgleichung mit der Ziel-Lösung ``x_val``.

    Die Koeffizienten werden so gewählt, dass bei der Kreuzmultiplikation
    keine \(x^2\)-Terme entstehen und ``x_val`` tatsächlich eine Lösung ist.
    """

    def fmt_lin(px: Fraction, q: Fraction) -> str:
        parts: List[str] = []
        if px != 0:
            parts.append(format_coeff_times_x(px))
        if q != 0:
            parts.append(fmt_term(q, with_sign=(len(parts) > 0)))
        return tidy_spaces(" ".join(parts)) if parts else "0"

    while True:
        # Proportionale Koeffizienten für x, damit sich x^2 nach
        # Kreuzmultiplikation wegkürzt: u1*w2 == u2*w1
        u1 = Fraction(rnd_nonzero(-6, 6))
        w1 = Fraction(rnd_nonzero(-6, 6))
        k = Fraction(rnd_nonzero(-5, 5))
        u2 = u1 * k
        w2 = w1 * k

        v1 = Fraction(randint(-10, 10))
        z1 = Fraction(randint(-10, 10))
        if w1 * x_val + z1 == 0:
            continue

        z2 = Fraction(randint(-10, 10))
        if w2 * x_val + z2 == 0:
            continue

        denom = w1 * x_val + z1
        if denom == 0:
            continue

        # Bestimme v2 so, dass x_val Lösung wird
        v2 = ((u1 * z2 + v1 * w2 - u2 * z1) * x_val + v1 * z2) / denom

        # Kanonische Form A x + B = 0
        bL = u1 * z2 + v1 * w2
        bR = u2 * z1 + v2 * w1
        cL = v1 * z2
        cR = v2 * z1
        A = bL - bR
        B = cL - cR
        if A == 0:
            continue

        canon = LinearEquation(A, B, Fraction(0), Fraction(0))

        left_text = f"({fmt_lin(u1, v1)})/({fmt_lin(w1, z1)})"
        right_text = f"({fmt_lin(u2, v2)})/({fmt_lin(w2, z2)})"
        txt = f"{left_text} = {right_text}"
        return DisplayEquation(text=txt, canonical=canon, kind="brueche")

# ---------------------- Generierung pro Kategorie -----------------------------

def gen_sehr_leicht(n: int) -> List[DisplayEquation]:
    out = []
    for _ in range(n):
        x = pick_target_x(CONFIG["ziel_mehr_brueche"])
        base = build_base_equation(x)
        eq = to_very_easy_form(base)
        out.append(DisplayEquation(eq.as_string(), eq, "sehr_leicht"))
    return out

def gen_leicht(n: int) -> List[DisplayEquation]:
    out = []
    for _ in range(n):
        x = pick_target_x(CONFIG["ziel_mehr_brueche"])
        eq = to_easy_both_sides(build_base_equation(x))
        out.append(DisplayEquation(eq.as_string(), eq, "leicht"))
    return out

def gen_mittel(n: int) -> List[DisplayEquation]:
    out = []
    for _ in range(n):
        x = pick_target_x(CONFIG["ziel_mehr_brueche"])
        out.append(to_medium_with_parenthesis(x))
    return out

def gen_schwer(n: int) -> List[DisplayEquation]:
    out = []
    for _ in range(n):
        x = pick_target_x(CONFIG["ziel_mehr_brueche"])
        out.append(to_harder_mixed_parenthesis(x))
    return out

def gen_brueche(n: int) -> List[DisplayEquation]:
    out = []
    for _ in range(n):
        x = pick_target_x(CONFIG["ziel_mehr_brueche"])
        out.append(to_fraction_equation(x))
    return out

# ---------------------- Lösungsschritte erzeugen ------------------------------

def solution_steps(eq: LinearEquation) -> List[str]:
    """
    Erzeuge einen klaren, knappen Rechenweg (auf Deutsch):
      1) Ausmultiplizieren / Klammern entfernen (falls nötig) – hier bereits kanonisch
      2) x-Terme auf eine Seite, Zahlen auf die andere
      3) Zusammenfassen und durch Koeffizient teilen
    """
    steps = []
    # Start: ax + b = cx + d
    a, b, c, d = eq.a, eq.b, eq.c, eq.d
    steps.append(f"Ausgangsgleichung: {eq.as_string()}")

    # x-Terme auf links: (a-c)x + b = d
    A = a - c
    steps.append(f"Bringe x-Terme auf die linke Seite: ({format_frac(a)} - {format_frac(c)})x + {format_frac(b)} = {format_frac(d)}")

    # Zahlen auf rechts: (a-c)x = d - b
    B = d - b
    steps.append(f"Bringe konstante Terme auf die rechte Seite: ({format_frac(A)})x = {format_frac(d)} - {format_frac(b)} = {format_frac(B)}")

    if A == 0:
        steps.append("Spezialfall: Kein eindeutiges x (entfällt hier, da Generator das vermeidet).")
        return steps

    # Division
    x_val = B / A
    steps.append(f"Teile durch {format_frac(A)}: x = {format_frac(B)} / {format_frac(A)} = {format_frac(x_val)}")

    # Optional: gemischte Zahl
    if x_val.denominator != 1:
        ganz = x_val.numerator // x_val.denominator
        rest = abs(x_val.numerator) % x_val.denominator
        if ganz != 0 and rest != 0:
            mix = f"{ganz} {rest}/{x_val.denominator}"
            steps.append(f"(Als gemischte Zahl: x = {mix})")
    return steps

# ---------------------- DOCX Ausgabe -----------------------------------------

def add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(CONFIG["schriftgroesse_punkte"] + 2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_equation_par(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}  {text}")
    run.font.size = Pt(CONFIG["schriftgroesse_punkte"])
    # Leerraum
    for _ in range(CONFIG["leerzeilen_zwischen_aufgaben"]):
        doc.add_paragraph(" ")

def write_arbeitsblatt(eqs: List[DisplayEquation], path: str):
    doc = Document()
    add_heading(doc, "Lineare Gleichungen – Arbeitsblatt")
    doc.add_paragraph(" ")
    # Nach Kategorien gruppieren, Reihenfolge fest, aber Aufgaben gemischt innerhalb
    order = ["sehr_leicht", "leicht", "mittel", "schwer", "brueche"]
    for kind in order:
        block = [e for e in eqs if e.kind == kind]
        if not block:
            continue
        add_heading(doc, kind.replace("_", " ").title())
        for i, deq in enumerate(block, 1):
            add_equation_par(doc, f"({i})", deq.text)
        doc.add_paragraph(" ")
    doc.save(path)

def write_loesungen(eqs: List[DisplayEquation], path: str):
    doc = Document()
    add_heading(doc, "Lineare Gleichungen – Lösungen & Rechenweg")
    doc.add_paragraph(" ")
    order = ["sehr_leicht", "leicht", "mittel", "schwer", "brueche"]
    for kind in order:
        block = [e for e in eqs if e.kind == kind]
        if not block:
            continue
        add_heading(doc, kind.replace("_", " ").title())
        for i, deq in enumerate(block, 1):
            # Kanonische Form zeigen (für Nachvollziehbarkeit)
            doc.add_paragraph(f"Aufgabe ({i}): {deq.text}")
            # Lösungsschritte aus kanonischer ax+b=cx+d-Form:
            for step in solution_steps(deq.canonical):
                p = doc.add_paragraph("• " + step)
                p.runs[0].font.size = Pt(CONFIG["schriftgroesse_punkte"])
            doc.add_paragraph(" ")
        doc.add_paragraph(" ")
    doc.save(path)

# ---------------------- Main --------------------------------------------------

def generate_all() -> Tuple[List[DisplayEquation], str, str]:
    tasks = []
    tasks += gen_sehr_leicht(CONFIG["counts"]["sehr_leicht"])
    tasks += gen_leicht(CONFIG["counts"]["leicht"])
    tasks += gen_mittel(CONFIG["counts"]["mittel"])
    tasks += gen_schwer(CONFIG["counts"]["schwer"])
    tasks += gen_brueche(CONFIG["counts"]["brueche"])

    # Option: durchmischen innerhalb der Kategorie-Reihenfolge belassen wir in DOCX.
    # (Arbeitsblatt/Lösungen gruppieren nach Kategorie)

    write_arbeitsblatt(tasks, CONFIG["arbeitsblatt_datei"])
    write_loesungen(tasks, CONFIG["loesungen_datei"])
    return tasks, CONFIG["arbeitsblatt_datei"], CONFIG["loesungen_datei"]

if __name__ == "__main__":
    _, ab, lb = generate_all()
    print(f"Fertig.\nArbeitsblatt: {ab}\nLösungen:    {lb}")
